from peewee import *
from pandas import *
from faker import Faker
from datetime import datetime
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import *
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from openpyxl import *
from reportlab.lib import colors
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.styles import getSampleStyleSheet
from docx.enum.table import WD_ALIGN_VERTICAL

# Игнор ошибки стилей таблиц
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="docx")

fake = Faker()

# Создание структуры таблицы
table_data = [
    {
        'Дата': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'Слабость, утомляемость': fake.random_int(min=0, max=3),
        'Рвота': fake.random_int(min=0, max=1),
        'Уменьшение диуреза': fake.random_int(min=0, max=1),
        'Ухудшение памяти': fake.random_int(min=0, max=1),
        'Нарушение моторных функций': fake.random_int(min=0, max=1),
        'Болевой синдром': fake.random_int(min=0, max=3),
        'Хрипы': fake.random_int(min=0, max=1),
        'Бронхоспазм': fake.random_int(min=0, max=1),
        'Боль в левой части грудной клетки': fake.random_int(min=0, max=1),
        'Аритмия': fake.random_int(min=0, max=1),
        'Депрессия, тревога': fake.random_int(min=0, max=3),
        'Спутанность сознания (химический мозг)': fake.random_int(min=0, max=1),
        'Прилив жара к верхней части туловища': fake.random_int(min=0, max=1),
        'Нейродермит (сыпь, зуд)': fake.random_int(min=0, max=1),
        'Стоматит': fake.random_int(min=0, max=1),
        'Периферическая невропатия руки': fake.random_int(min=0, max=1),
        'Периферическая невропатия ноги': fake.random_int(min=0, max=1),
        'Мигрень': fake.random_int(min=0, max=3)
        
    } for _ in range(5)
]

# Преобразование данных в DataFrame pandas
data = {key: [] for key in table_data[0].keys()}
for row in table_data:
    for key, value in row.items():
        data[key].append(value)

df = DataFrame(data)

current_time = datetime.now().strftime("%Y_%m_%d_%H_%M_%S")


def export_xlsx():
    excel_file_path = f'Отчет_{current_time}.xlsx'
    df.to_excel(excel_file_path, index=False, engine='openpyxl')
    wb = load_workbook(excel_file_path)
    ws = wb.active

    for column in ws.columns:
        max_length = 0
        column = [cell for cell in column]
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 7)
        ws.column_dimensions[column[0].column_letter].width = adjusted_width

    wb.save(excel_file_path)
    print(f'Data has been exported to {excel_file_path}')


# Сохранение данных в Word
def export_docx():
    global word_file_path, word_creation_time
    word_creation_time = current_time
    word_file_path = f'Отчет_{word_creation_time}.docx'

    document = Document()

    table = document.add_table(rows=len(df.columns) + 1, cols=len(df) + 1)

    table.autofit = False
    table.style = 'TableGrid'


    for row_num, col_name in enumerate(["Дата"] + list(df.columns)):
        cell = table.cell(row_num - 1, 0)
        cell.text = col_name
        cell.paragraphs[0].runs[0].font.name = 'SitkaVF'
        cell.paragraphs[0].runs[0].font.bold = True


    for col_num, row_data in enumerate(df.itertuples(), start=1):
        for row_num, value in enumerate([current_time] + list(row_data[1:]), start=0):
            cell = table.cell(row_num - 1, col_num)
            cell.text = str(value)
            cell.paragraphs[0].runs[0].font.name = 'SitkaVF'


    table.columns[0].width = Pt(100)

    for col_num in range(1, len(df) + 1):
        table.columns[col_num].width = Pt(70)

    # Выравнивание ячеек
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    document.save(word_file_path)
    print(f'Data has been exported to {word_file_path}')


# Сохранение данных в PDF
def export_pdf():
    styles = getSampleStyleSheet()  # дефолтовые стили
    
    styles['Normal'].fontName = 'SitkaVF'
    styles['Heading1'].fontName = 'SitkaVF'
    styles['Normal'].fontSize = 15
    styles['Heading1'].fontSize = 15
    pdfmetrics.registerFont(TTFont('SitkaVF', 'C:/Windows/Fonts/SitkaVF.ttf', 'UTF-8'))

    pdf_file_path = f"Отчет_{current_time}.pdf"
    doc = SimpleDocTemplate(pdf_file_path, pagesize=A2)

    story = []
    col_names = list(df.columns)

    data = []

    for col_num, col_name in enumerate(col_names):
        data.append([col_name] + [str(value) for value in df[col_name]])

    col_widths = [None] * len(data[0])  # None - авто определение ширины столбцов

    # Создание таблицы в ReportLab
    pdf_table = Table(data, colWidths=col_widths)
    style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.grey),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'SitkaVF'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 15),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('SPLITBYROW', (0, 0), (-1, -1)),  # Растягивать по высоте
    ])

    pdf_table.setStyle(style)
    story.append(pdf_table)

    doc.build(story)
    print(f'Data has been exported to {pdf_file_path}')

